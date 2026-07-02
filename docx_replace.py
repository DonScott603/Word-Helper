"""Core find-and-replace engine for .docx files.

Word's built-in Find & Replace dialog caps the search/replace strings at 255
characters. This module works directly on the document XML via python-docx, so
there is no length limit. It also handles text that Word has split across
multiple "runs" (which happens constantly due to spell-check marks, formatting
boundaries, etc.), replacing across run boundaries while preserving the
formatting of the run where each match begins.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from docx import Document


@dataclass
class ReplaceResult:
    """Outcome of processing a single document."""
    path: str
    replacements: int = 0
    ok: bool = True
    error: str = ""
    locations: dict = field(default_factory=dict)  # e.g. {"body": 3, "headers": 1}


def _iter_block_paragraphs(container):
    """Yield every paragraph in a container (document, cell, header, footer),
    recursing into tables and nested tables."""
    for paragraph in container.paragraphs:
        yield paragraph
    for table in container.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_block_paragraphs(cell)


def _iter_document_regions(document):
    """Yield (region_name, paragraph) for every paragraph in the document,
    including headers and footers of every section."""
    for paragraph in _iter_block_paragraphs(document):
        yield "body", paragraph

    for section in document.sections:
        headers = (
            section.header,
            section.first_page_header,
            section.even_page_header,
        )
        footers = (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        )
        for header in headers:
            if header is None:
                continue
            for paragraph in _iter_block_paragraphs(header):
                yield "headers", paragraph
        for footer in footers:
            if footer is None:
                continue
            for paragraph in _iter_block_paragraphs(footer):
                yield "footers", paragraph


def replace_in_paragraph(paragraph, find, replace, match_case=True):
    """Replace all occurrences of ``find`` with ``replace`` inside one paragraph.

    Works across run boundaries. The replacement text inherits the formatting of
    the run in which the match starts. Returns the number of replacements made.
    """
    runs = paragraph.runs
    if not runs or not find:
        return 0

    count = 0
    # Reprocess from scratch after each replacement, because editing run text
    # shifts every subsequent offset. Paragraphs are short, so this is cheap.
    while True:
        # Build the concatenated text plus a map from global offset -> run.
        run_texts = [r.text for r in runs]
        full_text = "".join(run_texts)

        haystack = full_text if match_case else full_text.lower()
        needle = find if match_case else find.lower()

        pos = haystack.find(needle)
        if pos == -1:
            break

        start = pos
        end = pos + len(find)

        # Locate the run containing the start and the run containing the last
        # matched character, tracking each run's starting offset.
        offsets = []
        acc = 0
        for text in run_texts:
            offsets.append(acc)
            acc += len(text)

        start_run = end_run = None
        for i, off in enumerate(offsets):
            run_end = off + len(run_texts[i])
            if start_run is None and start < run_end:
                start_run = i
            if off <= end - 1 < run_end:
                end_run = i
                break
        if start_run is None:
            break
        if end_run is None:
            end_run = len(runs) - 1

        start_off = offsets[start_run]
        end_off = offsets[end_run]
        prefix = run_texts[start_run][: start - start_off]
        suffix = run_texts[end_run][end - end_off:]

        if start_run == end_run:
            runs[start_run].text = prefix + replace + suffix
        else:
            runs[start_run].text = prefix + replace
            for i in range(start_run + 1, end_run):
                runs[i].text = ""
            runs[end_run].text = suffix

        count += 1

    return count


def replace_in_document(path, find, replace, match_case=True, scope="everywhere"):
    """Open ``path``, replace ``find`` with ``replace`` everywhere in scope, and
    return a ReplaceResult. Does NOT save; the caller decides how to persist.

    ``scope`` is one of: "everywhere", "body", "body_paragraphs".
    Returns (document, ReplaceResult). The document is returned so the caller can
    save it (in place, to a copy, etc.).
    """
    result = ReplaceResult(path=path)
    try:
        document = Document(path)
    except Exception as exc:  # pragma: no cover - surfaced in the GUI log
        result.ok = False
        result.error = str(exc)
        return None, result

    for region, paragraph in _iter_document_regions(document):
        if scope == "body" and region != "body":
            continue
        made = replace_in_paragraph(paragraph, find, replace, match_case=match_case)
        if made:
            result.replacements += made
            result.locations[region] = result.locations.get(region, 0) + made

    return document, result
