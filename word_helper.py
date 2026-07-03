"""Word Helper — a small dark-themed GUI for batch operations on .docx files.

Tabs:
  * Find & Replace  — batch text find/replace with no 255-character limit.
  * Replace Image   — swap an embedded image (old file -> new file) across many
                      documents at once.

The layout is modular: each feature is a build_*_tab method, and the file list
(FileSelector) and image pickers (ImagePicker) are reusable widgets.
"""

from __future__ import annotations

import io
import os
import queue
import shutil
import threading
from tkinter import filedialog

import customtkinter as ctk
from PIL import Image as PILImage

from docx_replace import replace_in_document
from docx_image_replace import replace_images_in_document, read_image_info
from docx_format import apply_formatting_in_document, FormatSpec
from docx_pages import extract_pages, add_pages, count_pages

APP_NAME = "Word Helper"
ACCENT = "#3d7eff"
ACCENT_HOVER = "#2f66d6"
GREY = "#8a8f98"
CARD = "#2b2d31"

TRISTATE = {"Leave": None, "On": True, "Off": False}

ctk.set_appearance_mode("dark")
ctk.set_default_color_theme("dark-blue")

DOCX_TYPES = [("Word documents", "*.docx"), ("All files", "*.*")]
IMAGE_TYPES = [
    ("Images", "*.png *.jpg *.jpeg *.gif *.bmp *.tif *.tiff *.emf *.wmf"),
    ("All files", "*.*"),
]


# --------------------------------------------------------------------- widgets
class FileSelector(ctk.CTkFrame):
    """A titled, scrollable list of selected .docx files with add/clear/remove."""

    def __init__(self, master, title="Word files", height=140):
        super().__init__(master, fg_color="transparent")
        self.files: list[str] = []

        bar = ctk.CTkFrame(self, fg_color="transparent")
        bar.pack(fill="x")
        ctk.CTkLabel(
            bar, text=title, font=ctk.CTkFont(size=15, weight="bold")
        ).pack(side="left")
        ctk.CTkButton(
            bar, text="Clear", width=70, fg_color="#3a3d44",
            hover_color="#4a4e57", command=self.clear,
        ).pack(side="right")
        ctk.CTkButton(
            bar, text="+ Add files", width=110, fg_color=ACCENT,
            hover_color=ACCENT_HOVER, command=self.add,
        ).pack(side="right", padx=(0, 8))

        self._list = ctk.CTkScrollableFrame(self, height=height, label_text="")
        self._list.pack(fill="both", expand=True, pady=(6, 0))
        self._render()

    def add(self):
        for p in filedialog.askopenfilenames(
            title="Select Word documents", filetypes=DOCX_TYPES
        ):
            if p not in self.files:
                self.files.append(p)
        self._render()

    def clear(self):
        self.files.clear()
        self._render()

    def remove(self, path):
        if path in self.files:
            self.files.remove(path)
        self._render()

    def _render(self):
        for child in self._list.winfo_children():
            child.destroy()
        if not self.files:
            ctk.CTkLabel(
                self._list, text="No files selected.", text_color=GREY
            ).pack(anchor="w", padx=6, pady=6)
            return
        for path in self.files:
            row = ctk.CTkFrame(self._list, fg_color=CARD)
            row.pack(fill="x", pady=2, padx=2)
            ctk.CTkButton(
                row, text="✕", width=28, height=24, fg_color="#3a3d44",
                hover_color="#a13d3d", command=lambda p=path: self.remove(p),
            ).pack(side="right", padx=4, pady=2)
            ctk.CTkLabel(row, text=os.path.basename(path), anchor="w").pack(
                side="left", padx=8, pady=2
            )


class ImagePicker(ctk.CTkFrame):
    """A card that lets the user choose an image file and shows a thumbnail +
    format/dimension info. ``bytes`` holds the chosen file's raw data."""

    def __init__(self, master, title):
        super().__init__(master, fg_color=CARD, corner_radius=8)
        self.path: str | None = None
        self.bytes: bytes | None = None
        self._imgref = None

        ctk.CTkLabel(
            self, text=title, font=ctk.CTkFont(size=14, weight="bold")
        ).pack(pady=(10, 6))
        self.preview = ctk.CTkLabel(
            self, text="no image", width=140, height=140,
            fg_color="#1f2124", corner_radius=6, text_color=GREY,
        )
        self.preview.pack(padx=12)
        self.info = ctk.CTkLabel(
            self, text="", text_color=GREY, font=ctk.CTkFont(size=12),
            justify="center",
        )
        self.info.pack(pady=(6, 4))
        ctk.CTkButton(
            self, text="Choose…", width=120, fg_color=ACCENT,
            hover_color=ACCENT_HOVER, command=self._choose,
        ).pack(pady=(0, 12))

    def valid(self) -> bool:
        return self.bytes is not None

    def _choose(self):
        path = filedialog.askopenfilename(title="Select image", filetypes=IMAGE_TYPES)
        if not path:
            return
        try:
            with open(path, "rb") as fh:
                data = fh.read()
        except Exception as exc:
            self.info.configure(text=f"could not read file:\n{exc}")
            return
        self.path = path
        self.bytes = data

        try:
            info = read_image_info(data)
            self.info.configure(
                text=f"{os.path.basename(path)}\n"
                     f"{info.px_width}×{info.px_height} px · {info.ext.upper()}"
            )
        except Exception:
            self.info.configure(text=f"{os.path.basename(path)}\n(format not recognised)")

        try:
            im = PILImage.open(io.BytesIO(data))
            im.thumbnail((130, 130))
            self._imgref = ctk.CTkImage(light_image=im, dark_image=im, size=im.size)
            self.preview.configure(image=self._imgref, text="")
        except Exception:
            self._imgref = None
            self.preview.configure(image="", text="(no preview)")


class FilePick(ctk.CTkFrame):
    """A single-file chooser: label + button + selected-name/info line."""

    def __init__(self, master, label, filetypes=DOCX_TYPES, on_change=None):
        super().__init__(master, fg_color="transparent")
        self.path: str | None = None
        self.filetypes = filetypes
        self.on_change = on_change
        ctk.CTkLabel(self, text=label, width=120, anchor="w").pack(side="left")
        ctk.CTkButton(
            self, text="Choose…", width=90, fg_color=ACCENT,
            hover_color=ACCENT_HOVER, command=self._pick,
        ).pack(side="left", padx=(0, 8))
        self.info = ctk.CTkLabel(self, text="(none)", text_color=GREY, anchor="w")
        self.info.pack(side="left", fill="x", expand=True)

    def _pick(self):
        p = filedialog.askopenfilename(title="Select document", filetypes=self.filetypes)
        if not p:
            return
        self.path = p
        self.info.configure(text=os.path.basename(p))
        if self.on_change:
            self.on_change(p, self.info)


# ------------------------------------------------------------------------- app
class WordHelperApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title(APP_NAME)
        self.geometry("980x860")
        self.minsize(820, 680)

        self._msg_queue: queue.Queue = queue.Queue()
        self._running = False

        self._build_header()

        self.tabview = ctk.CTkTabview(self, fg_color="transparent")
        self.tabview.pack(fill="both", expand=True, padx=16, pady=(0, 12))

        self.tabview.add("Find & Replace")
        self.build_find_replace_tab(self.tabview.tab("Find & Replace"))

        self.tabview.add("Replace Image")
        self.build_image_replace_tab(self.tabview.tab("Replace Image"))

        self.tabview.add("Formatting")
        self.build_formatting_tab(self.tabview.tab("Formatting"))

        self.tabview.add("Add / Extract Pages")
        self.build_pages_tab(self.tabview.tab("Add / Extract Pages"))

        self.after(100, self._drain_queue)

    # ------------------------------------------------------------------ header
    def _build_header(self):
        header = ctk.CTkFrame(self, fg_color="transparent")
        header.pack(fill="x", padx=16, pady=(14, 6))
        ctk.CTkLabel(
            header, text="Word Helper", font=ctk.CTkFont(size=24, weight="bold")
        ).pack(side="left")
        ctk.CTkLabel(
            header, text="Batch tools for .docx files",
            font=ctk.CTkFont(size=13), text_color=GREY,
        ).pack(side="left", padx=(12, 0), pady=(6, 0))

    # ------------------------------------------------ shared run-panel builder
    def _build_run_panel(self, tab, row, button_text, command):
        """Create a run button + progress bar (one row) and a log box (next
        row, expanding). Returns (run_btn, progress, log_box)."""
        run_bar = ctk.CTkFrame(tab, fg_color="transparent")
        run_bar.grid(row=row, column=0, sticky="ew", pady=(6, 6))
        run_btn = ctk.CTkButton(
            run_bar, text=button_text, width=170, height=40,
            font=ctk.CTkFont(size=15, weight="bold"),
            fg_color=ACCENT, hover_color=ACCENT_HOVER, command=command,
        )
        run_btn.pack(side="left")
        progress = ctk.CTkProgressBar(run_bar)
        progress.set(0)
        progress.pack(side="left", fill="x", expand=True, padx=(14, 0))

        log_box = ctk.CTkTextbox(tab, wrap="word", font=ctk.CTkFont(size=12))
        log_box.grid(row=row + 1, column=0, sticky="nsew", pady=(6, 4))
        log_box.configure(state="disabled")
        return run_btn, progress, log_box

    # -------------------------------------------------------- find & replace UI
    def build_find_replace_tab(self, tab):
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(0, weight=1)  # file list grows
        tab.grid_rowconfigure(5, weight=1)  # log grows

        self.fr_files = FileSelector(tab)
        self.fr_files.grid(row=0, column=0, sticky="nsew", pady=(8, 10))

        io_frame = ctk.CTkFrame(tab, fg_color="transparent")
        io_frame.grid(row=1, column=0, sticky="ew")
        io_frame.grid_columnconfigure(0, weight=1)
        io_frame.grid_columnconfigure(1, weight=1)
        ctk.CTkLabel(io_frame, text="Find").grid(row=0, column=0, sticky="w")
        ctk.CTkLabel(io_frame, text="Replace with").grid(
            row=0, column=1, sticky="w", padx=(10, 0)
        )
        self.fr_find = ctk.CTkTextbox(io_frame, height=110, wrap="word")
        self.fr_find.grid(row=1, column=0, sticky="nsew", pady=(2, 0))
        self.fr_replace = ctk.CTkTextbox(io_frame, height=110, wrap="word")
        self.fr_replace.grid(row=1, column=1, sticky="nsew", padx=(10, 0), pady=(2, 0))

        opts = ctk.CTkFrame(tab, fg_color="transparent")
        opts.grid(row=2, column=0, sticky="ew", pady=(10, 4))
        self.fr_match_case = ctk.CTkSwitch(opts, text="Match case")
        self.fr_match_case.pack(side="left")
        self.fr_scope = ctk.StringVar(value="everywhere")
        ctk.CTkLabel(opts, text="Scope:").pack(side="left", padx=(20, 6))
        ctk.CTkOptionMenu(
            opts, width=180, values=["everywhere", "body"], variable=self.fr_scope,
            fg_color="#3a3d44", button_color="#3a3d44", button_hover_color="#4a4e57",
        ).pack(side="left")
        ctk.CTkLabel(
            opts, text="Originals are backed up to .bak before overwriting.",
            text_color=GREY, font=ctk.CTkFont(size=12),
        ).pack(side="right")

        self.fr_run, self.fr_progress, self.fr_log = self._build_run_panel(
            tab, 4, "Run replace", self._start_text_run
        )
        self._log(self.fr_log,
                  "Add .docx files, enter your Find and Replace text, then Run.")

    # --------------------------------------------------------- replace image UI
    def build_image_replace_tab(self, tab):
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(2, weight=1)  # file list grows
        tab.grid_rowconfigure(5, weight=1)  # log grows

        pickers = ctk.CTkFrame(tab, fg_color="transparent")
        pickers.grid(row=0, column=0, sticky="ew", pady=(8, 6))
        pickers.grid_columnconfigure(0, weight=1)
        pickers.grid_columnconfigure(1, weight=0)
        pickers.grid_columnconfigure(2, weight=1)
        self.im_old = ImagePicker(pickers, "Old image (find)")
        self.im_old.grid(row=0, column=0, sticky="ew", padx=(0, 6))
        ctk.CTkLabel(pickers, text="→", font=ctk.CTkFont(size=28)).grid(
            row=0, column=1, padx=8
        )
        self.im_new = ImagePicker(pickers, "New image (replace with)")
        self.im_new.grid(row=0, column=2, sticky="ew", padx=(6, 0))

        opts = ctk.CTkFrame(tab, fg_color="transparent")
        opts.grid(row=1, column=0, sticky="ew", pady=(8, 4))
        self.im_dims = ctk.CTkSwitch(
            opts, text="Also match by size if not an exact match"
        )
        self.im_dims.pack(side="left")
        ctk.CTkLabel(
            opts,
            text="Matches images identical to the old file; new image keeps the "
                 "original's size. Originals backed up to .bak.",
            text_color=GREY, font=ctk.CTkFont(size=12), wraplength=380, justify="right",
        ).pack(side="right")

        self.im_files = FileSelector(tab)
        self.im_files.grid(row=2, column=0, sticky="nsew", pady=(8, 10))

        self.im_run, self.im_progress, self.im_log = self._build_run_panel(
            tab, 4, "Replace image", self._start_image_run
        )
        self._log(self.im_log,
                  "Choose the old and new image, add .docx files, then Replace.")

    # -------------------------------------------------------------- formatting UI
    def _tristate(self, parent, label, r, c):
        cell = ctk.CTkFrame(parent, fg_color="transparent")
        cell.grid(row=r, column=c, padx=(0, 18), pady=2, sticky="w")
        ctk.CTkLabel(cell, text=label, font=ctk.CTkFont(size=12)).pack(anchor="w")
        seg = ctk.CTkSegmentedButton(cell, values=["Leave", "On", "Off"])
        seg.set("Leave")
        seg.pack()
        return seg

    def build_formatting_tab(self, tab):
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(0, weight=1)  # file list grows
        tab.grid_rowconfigure(6, weight=1)  # log grows

        self.fm_files = FileSelector(tab)
        self.fm_files.grid(row=0, column=0, sticky="nsew", pady=(8, 10))

        ctk.CTkLabel(
            tab, text="Find text to format",
            font=ctk.CTkFont(size=15, weight="bold"),
        ).grid(row=1, column=0, sticky="w")
        self.fm_find = ctk.CTkTextbox(tab, height=70, wrap="word")
        self.fm_find.grid(row=2, column=0, sticky="ew", pady=(2, 0))

        opts = ctk.CTkFrame(tab, fg_color="transparent")
        opts.grid(row=3, column=0, sticky="ew", pady=(10, 4))
        self.fm_match_case = ctk.CTkSwitch(opts, text="Match case")
        self.fm_match_case.pack(side="left")
        self.fm_scope = ctk.StringVar(value="everywhere")
        ctk.CTkLabel(opts, text="Scope:").pack(side="left", padx=(20, 6))
        ctk.CTkOptionMenu(
            opts, width=180, values=["everywhere", "body"], variable=self.fm_scope,
            fg_color="#3a3d44", button_color="#3a3d44", button_hover_color="#4a4e57",
        ).pack(side="left")
        ctk.CTkLabel(
            opts, text="Only the options you set are applied — 'Leave' means unchanged.",
            text_color=GREY, font=ctk.CTkFont(size=12),
        ).pack(side="right")

        fmt = ctk.CTkFrame(tab, fg_color=CARD, corner_radius=8)
        fmt.grid(row=4, column=0, sticky="ew", pady=(6, 4))
        inner = ctk.CTkFrame(fmt, fg_color="transparent")
        inner.pack(fill="x", padx=14, pady=12)

        self.fm_bold = self._tristate(inner, "Bold", 0, 0)
        self.fm_italic = self._tristate(inner, "Italic", 0, 1)
        self.fm_underline = self._tristate(inner, "Underline", 0, 2)
        self.fm_strike = self._tristate(inner, "Strikethrough", 0, 3)

        line = ctk.CTkFrame(inner, fg_color="transparent")
        line.grid(row=1, column=0, columnspan=4, sticky="ew", pady=(14, 0))
        self.fm_set_font = ctk.CTkCheckBox(line, text="Font", width=70)
        self.fm_set_font.pack(side="left")
        self.fm_font = ctk.CTkEntry(line, width=150, placeholder_text="e.g. Arial")
        self.fm_font.pack(side="left", padx=(4, 20))
        self.fm_set_size = ctk.CTkCheckBox(line, text="Size (pt)", width=90)
        self.fm_set_size.pack(side="left")
        self.fm_size = ctk.CTkEntry(line, width=60, placeholder_text="12")
        self.fm_size.pack(side="left", padx=(4, 20))
        self.fm_set_color = ctk.CTkCheckBox(line, text="Color (hex)", width=100)
        self.fm_set_color.pack(side="left")
        self.fm_color = ctk.CTkEntry(line, width=90, placeholder_text="FF0000")
        self.fm_color.pack(side="left", padx=(4, 0))

        self.fm_run, self.fm_progress, self.fm_log = self._build_run_panel(
            tab, 5, "Apply formatting", self._start_format_run
        )
        self._log(self.fm_log,
                  "Add files, enter the text to find, choose formatting, then Apply.")

    # ------------------------------------------------------- run: formatting tab
    def _start_format_run(self):
        if self._running:
            return
        find = self.fm_find.get("1.0", "end-1c")
        if not self.fm_files.files:
            self._log(self.fm_log, "⚠ No files selected.")
            return
        if not find:
            self._log(self.fm_log, "⚠ The Find box is empty.")
            return

        spec = FormatSpec(
            bold=TRISTATE[self.fm_bold.get()],
            italic=TRISTATE[self.fm_italic.get()],
            underline=TRISTATE[self.fm_underline.get()],
            strike=TRISTATE[self.fm_strike.get()],
        )
        if self.fm_set_font.get() and self.fm_font.get().strip():
            spec.name = self.fm_font.get().strip()
        if self.fm_set_size.get():
            try:
                spec.size_pt = float(self.fm_size.get().strip())
                if spec.size_pt <= 0:
                    raise ValueError
            except ValueError:
                self._log(self.fm_log, "⚠ Size must be a positive number (points).")
                return
        if self.fm_set_color.get():
            hexv = self.fm_color.get().strip().lstrip("#").upper()
            if len(hexv) != 6 or any(c not in "0123456789ABCDEF" for c in hexv):
                self._log(self.fm_log, "⚠ Color must be 6 hex digits, e.g. FF0000.")
                return
            spec.color = hexv

        if spec.is_empty():
            self._log(self.fm_log,
                      "⚠ No formatting selected — set at least one option.")
            return

        self._set_running(self.fm_run, True, "Working…")
        self.fm_progress.set(0)
        self._log(self.fm_log, f"\n── Starting: {len(self.fm_files.files)} file(s) ──")
        threading.Thread(
            target=self._format_worker, daemon=True,
            kwargs=dict(
                files=list(self.fm_files.files), find=find, spec=spec,
                match_case=bool(self.fm_match_case.get()), scope=self.fm_scope.get(),
            ),
        ).start()

    def _format_worker(self, files, find, spec, match_case, scope):
        total = len(files)
        grand = changed = 0
        for i, path in enumerate(files, start=1):
            name = os.path.basename(path)
            document, result = apply_formatting_in_document(
                path, find, spec, match_case=match_case, scope=scope
            )
            if not result.ok:
                self._logq(self.fm_log, f"✖ {name}: {result.error}")
            elif result.matches == 0:
                self._logq(self.fm_log, f"•  {name}: no matches")
            else:
                loc = ", ".join(f"{k}: {v}" for k, v in result.locations.items())
                saved = self._save_with_backup(
                    path, document, self.fm_log, name,
                    f"{result.matches} match(es) formatted ({loc})",
                )
                if saved:
                    grand += result.matches
                    changed += 1
            self._progq(self.fm_progress, i / total)
        self._logq(self.fm_log,
                   f"── Done: {grand} match(es) formatted across {changed} file(s). ──")
        self._doneq(self.fm_run, "Apply formatting")

    # -------------------------------------------------------------- pages UI
    def _pages_show_count(self, path, label):
        """Count pages in a background Word instance and append to the label."""
        base = os.path.basename(path)
        label.configure(text=f"{base}  ·  counting…")

        def work():
            try:
                n = count_pages(path)
                self._ui(lambda: label.configure(text=f"{base}  ·  {n} pages"))
            except Exception:
                self._ui(lambda: label.configure(text=base))
        threading.Thread(target=work, daemon=True).start()

    def _choose_save(self, attr, label):
        p = filedialog.asksaveasfilename(
            title="Save result as", defaultextension=".docx", filetypes=DOCX_TYPES
        )
        if p:
            setattr(self, attr, p)
            label.configure(text=os.path.basename(p))

    def build_pages_tab(self, tab):
        tab.grid_columnconfigure(0, weight=1)
        tab.grid_rowconfigure(3, weight=1)  # log grows

        note = ctk.CTkLabel(
            tab,
            text="Uses Microsoft Word to work with real pages — Word must be "
                 "installed. Operations open Word briefly, so they take a moment.",
            text_color=GREY, font=ctk.CTkFont(size=12), justify="left",
        )
        note.grid(row=0, column=0, sticky="w", pady=(8, 6))

        self.pg_mode = ctk.CTkSegmentedButton(
            tab, values=["Extract pages", "Add pages"], command=self._pg_switch_mode
        )
        self.pg_mode.set("Extract pages")
        self.pg_mode.grid(row=1, column=0, sticky="w", pady=(0, 8))

        container = ctk.CTkFrame(tab, fg_color=CARD, corner_radius=8)
        container.grid(row=2, column=0, sticky="ew")
        container.grid_columnconfigure(0, weight=1)

        # --- Extract frame ---
        self._pg_extract = ext = ctk.CTkFrame(container, fg_color="transparent")
        ext.grid(row=0, column=0, sticky="ew", padx=14, pady=12)
        ext.grid_columnconfigure(0, weight=1)
        self.pg_ex_src = FilePick(ext, "Source document", on_change=self._pages_show_count)
        self.pg_ex_src.grid(row=0, column=0, sticky="ew", pady=4)
        rng = ctk.CTkFrame(ext, fg_color="transparent")
        rng.grid(row=1, column=0, sticky="w", pady=4)
        ctk.CTkLabel(rng, text="Pages   from").pack(side="left")
        self.pg_ex_from = ctk.CTkEntry(rng, width=60)
        self.pg_ex_from.insert(0, "1")
        self.pg_ex_from.pack(side="left", padx=6)
        ctk.CTkLabel(rng, text="to").pack(side="left")
        self.pg_ex_to = ctk.CTkEntry(rng, width=60)
        self.pg_ex_to.insert(0, "1")
        self.pg_ex_to.pack(side="left", padx=6)
        out = ctk.CTkFrame(ext, fg_color="transparent")
        out.grid(row=2, column=0, sticky="ew", pady=4)
        ctk.CTkLabel(out, text="Save extracted to", width=120, anchor="w").pack(side="left")
        self.pg_ex_out_label = ctk.CTkLabel(out, text="(none)", text_color=GREY, anchor="w")
        ctk.CTkButton(
            out, text="Choose…", width=90, fg_color=ACCENT, hover_color=ACCENT_HOVER,
            command=lambda: self._choose_save("pg_ex_out_path", self.pg_ex_out_label),
        ).pack(side="left", padx=(0, 8))
        self.pg_ex_out_label.pack(side="left", fill="x", expand=True)
        self.pg_ex_remove = ctk.CTkCheckBox(
            ext, text="Also remove these pages from the source (.bak backup)"
        )
        self.pg_ex_remove.grid(row=3, column=0, sticky="w", pady=(8, 2))

        # --- Add frame ---
        self._pg_add = add = ctk.CTkFrame(container, fg_color="transparent")
        add.grid(row=0, column=0, sticky="ew", padx=14, pady=12)
        add.grid_columnconfigure(0, weight=1)
        self.pg_add_tgt = FilePick(add, "Add into (target)", on_change=self._pages_show_count)
        self.pg_add_tgt.grid(row=0, column=0, sticky="ew", pady=4)
        self.pg_add_src = FilePick(add, "Pages from (source)", on_change=self._pages_show_count)
        self.pg_add_src.grid(row=1, column=0, sticky="ew", pady=4)

        srow = ctk.CTkFrame(add, fg_color="transparent")
        srow.grid(row=2, column=0, sticky="w", pady=4)
        ctk.CTkLabel(srow, text="Source pages:").pack(side="left", padx=(0, 6))
        self.pg_add_srcmode = ctk.CTkSegmentedButton(
            srow, values=["All", "Range"], command=lambda _: self._pg_toggle_srcrange()
        )
        self.pg_add_srcmode.set("All")
        self.pg_add_srcmode.pack(side="left")
        self.pg_add_sfrom = ctk.CTkEntry(srow, width=54)
        self.pg_add_sfrom.insert(0, "1")
        self.pg_add_sto = ctk.CTkEntry(srow, width=54)
        self.pg_add_sto.insert(0, "1")

        prow = ctk.CTkFrame(add, fg_color="transparent")
        prow.grid(row=3, column=0, sticky="w", pady=4)
        ctk.CTkLabel(prow, text="Insert:").pack(side="left", padx=(0, 6))
        self.pg_add_pos = ctk.CTkSegmentedButton(
            prow, values=["At end", "After page"],
            command=lambda _: self._pg_toggle_afterpage(),
        )
        self.pg_add_pos.set("At end")
        self.pg_add_pos.pack(side="left")
        self.pg_add_posnum = ctk.CTkEntry(prow, width=54)
        self.pg_add_posnum.insert(0, "1")

        orow = ctk.CTkFrame(add, fg_color="transparent")
        orow.grid(row=4, column=0, sticky="ew", pady=4)
        ctk.CTkLabel(orow, text="Output:").pack(side="left", padx=(0, 6))
        self.pg_add_outmode = ctk.CTkSegmentedButton(
            orow, values=["New file", "Overwrite target"],
            command=lambda _: self._pg_toggle_output(),
        )
        self.pg_add_outmode.set("New file")
        self.pg_add_outmode.pack(side="left")
        self.pg_add_out_btn = ctk.CTkButton(
            orow, text="Choose…", width=90, fg_color=ACCENT, hover_color=ACCENT_HOVER,
            command=lambda: self._choose_save("pg_add_out_path", self.pg_add_out_label),
        )
        self.pg_add_out_btn.pack(side="left", padx=(8, 8))
        self.pg_add_out_label = ctk.CTkLabel(orow, text="(none)", text_color=GREY, anchor="w")
        self.pg_add_out_label.pack(side="left", fill="x", expand=True)

        self.pg_add_hdrnote = ctk.CTkLabel(
            add,
            text="“At end” keeps the added document’s own headers/footers. "
                 "“After page” inserts mid-document, where the added pages take "
                 "on the target’s headers (a Word limitation).",
            text_color=GREY, font=ctk.CTkFont(size=12), wraplength=760, justify="left",
        )
        self.pg_add_hdrnote.grid(row=5, column=0, sticky="w", pady=(8, 0))

        self.pg_run, self.pg_progress, self.pg_log = self._build_run_panel(
            tab, 2 + 1, "Extract", self._start_pages_run
        )
        # run panel was gridded at rows 3/4; move log weight there
        tab.grid_rowconfigure(3, weight=0)
        tab.grid_rowconfigure(4, weight=1)

        self._pg_toggle_srcrange()
        self._pg_toggle_afterpage()
        self._pg_toggle_output()
        self._pg_switch_mode("Extract pages")
        self._log(self.pg_log, "Choose a mode, fill in the fields, then run.")

    def _pg_switch_mode(self, mode):
        if mode == "Extract pages":
            self._pg_add.grid_remove()
            self._pg_extract.grid()
            self.pg_run.configure(text="Extract")
        else:
            self._pg_extract.grid_remove()
            self._pg_add.grid()
            self.pg_run.configure(text="Add")

    def _pg_toggle_srcrange(self):
        if self.pg_add_srcmode.get() == "Range":
            self.pg_add_sfrom.pack(side="left", padx=(10, 4))
            self.pg_add_sto.pack(side="left", padx=4)
        else:
            self.pg_add_sfrom.pack_forget()
            self.pg_add_sto.pack_forget()

    def _pg_toggle_afterpage(self):
        if self.pg_add_pos.get() == "After page":
            self.pg_add_posnum.pack(side="left", padx=(8, 0))
        else:
            self.pg_add_posnum.pack_forget()

    def _pg_toggle_output(self):
        new_file = self.pg_add_outmode.get() == "New file"
        if new_file:
            self.pg_add_out_btn.configure(state="normal")
        else:
            self.pg_add_out_btn.configure(state="disabled")

    # ------------------------------------------------------------ run: pages
    def _parse_int(self, entry, name, log):
        try:
            v = int(entry.get().strip())
            if v < 1:
                raise ValueError
            return v
        except ValueError:
            self._log(log, f"⚠ {name} must be a whole number ≥ 1.")
            return None

    def _start_pages_run(self):
        if self._running:
            return
        if self.pg_mode.get() == "Extract pages":
            self._run_extract()
        else:
            self._run_add()

    def _run_extract(self):
        src = self.pg_ex_src.path
        out = getattr(self, "pg_ex_out_path", None)
        if not src:
            self._log(self.pg_log, "⚠ Choose a source document.")
            return
        if not out:
            self._log(self.pg_log, "⚠ Choose where to save the extracted pages.")
            return
        frm = self._parse_int(self.pg_ex_from, "From page", self.pg_log)
        to = self._parse_int(self.pg_ex_to, "To page", self.pg_log)
        if frm is None or to is None:
            return
        if to < frm:
            self._log(self.pg_log, "⚠ 'To' page must be ≥ 'From' page.")
            return
        remove = bool(self.pg_ex_remove.get())

        self._set_running(self.pg_run, True, "Working…")
        self._pg_busy(True)
        self._log(self.pg_log, f"\n── Extracting pages {frm}–{to} … ──")

        def work():
            result = extract_pages(src, out, frm, to, remove_from_source=remove)
            if result.ok:
                msg = (f"✔ Extracted {result.pages_affected} page(s) → "
                       f"{os.path.basename(result.output)}")
                if result.source_removed:
                    msg += f"; removed {result.source_removed} page(s) from source"
                self._logq(self.pg_log, msg)
            else:
                self._logq(self.pg_log, f"✖ {result.error}")
            self._pg_busyq(False)
            self._doneq(self.pg_run, "Extract")
        threading.Thread(target=work, daemon=True).start()

    def _run_add(self):
        target = self.pg_add_tgt.path
        source = self.pg_add_src.path
        if not target:
            self._log(self.pg_log, "⚠ Choose a target document to add into.")
            return
        if not source:
            self._log(self.pg_log, "⚠ Choose a source document to add from.")
            return

        insert_after = "end"
        if self.pg_add_pos.get() == "After page":
            page = self._parse_int(self.pg_add_posnum, "After page", self.pg_log)
            if page is None:
                return
            insert_after = page

        s_from = s_to = None
        if self.pg_add_srcmode.get() == "Range":
            s_from = self._parse_int(self.pg_add_sfrom, "Source from", self.pg_log)
            s_to = self._parse_int(self.pg_add_sto, "Source to", self.pg_log)
            if s_from is None or s_to is None:
                return
            if s_to < s_from:
                self._log(self.pg_log, "⚠ Source 'to' must be ≥ 'from'.")
                return

        dest = None
        if self.pg_add_outmode.get() == "New file":
            dest = getattr(self, "pg_add_out_path", None)
            if not dest:
                self._log(self.pg_log, "⚠ Choose an output file (or switch to Overwrite target).")
                return

        self._set_running(self.pg_run, True, "Working…")
        self._pg_busy(True)
        where = "at end" if insert_after == "end" else f"after page {insert_after}"
        self._log(self.pg_log, f"\n── Adding {os.path.basename(source)} {where} … ──")

        def work():
            result = add_pages(target, source, dest=dest, insert_after=insert_after,
                               src_start=s_from, src_end=s_to)
            if result.ok:
                self._logq(
                    self.pg_log,
                    f"✔ Added → {os.path.basename(result.output)} "
                    f"({result.pages_affected} page(s) total)",
                )
            else:
                self._logq(self.pg_log, f"✖ {result.error}")
            self._pg_busyq(False)
            self._doneq(self.pg_run, "Add")
        threading.Thread(target=work, daemon=True).start()

    def _pg_busy(self, on):
        if on:
            self.pg_progress.configure(mode="indeterminate")
            self.pg_progress.start()
        else:
            self.pg_progress.stop()
            self.pg_progress.configure(mode="determinate")
            self.pg_progress.set(1)

    def _pg_busyq(self, on):
        self._ui(lambda: self._pg_busy(on))

    # ----------------------------------------------------------- run: text tab
    def _start_text_run(self):
        if self._running:
            return
        find = self.fr_find.get("1.0", "end-1c")
        replace = self.fr_replace.get("1.0", "end-1c")
        if not self.fr_files.files:
            self._log(self.fr_log, "⚠ No files selected.")
            return
        if not find:
            self._log(self.fr_log, "⚠ The Find box is empty.")
            return

        self._set_running(self.fr_run, True, "Running…")
        self.fr_progress.set(0)
        self._log(self.fr_log, f"\n── Starting: {len(self.fr_files.files)} file(s) ──")
        threading.Thread(
            target=self._text_worker, daemon=True,
            kwargs=dict(
                files=list(self.fr_files.files), find=find, replace=replace,
                match_case=bool(self.fr_match_case.get()), scope=self.fr_scope.get(),
            ),
        ).start()

    def _text_worker(self, files, find, replace, match_case, scope):
        total = len(files)
        grand = changed = 0
        for i, path in enumerate(files, start=1):
            name = os.path.basename(path)
            document, result = replace_in_document(
                path, find, replace, match_case=match_case, scope=scope
            )
            if not result.ok:
                self._logq(self.fr_log, f"✖ {name}: {result.error}")
            elif result.replacements == 0:
                self._logq(self.fr_log, f"•  {name}: no matches")
            else:
                saved = self._save_with_backup(path, document, self.fr_log, name,
                                               f"{result.replacements} replaced")
                if saved:
                    loc = ", ".join(f"{k}: {v}" for k, v in result.locations.items())
                    self._logq(self.fr_log,
                               f"    ({loc})")
                    grand += result.replacements
                    changed += 1
            self._progq(self.fr_progress, i / total)
        self._logq(self.fr_log,
                   f"── Done: {grand} replacement(s) across {changed} file(s). ──")
        self._doneq(self.fr_run, "Run replace")

    # ---------------------------------------------------------- run: image tab
    def _start_image_run(self):
        if self._running:
            return
        if not self.im_old.valid() or not self.im_new.valid():
            self._log(self.im_log, "⚠ Choose both an old and a new image.")
            return
        if not self.im_files.files:
            self._log(self.im_log, "⚠ No files selected.")
            return

        self._set_running(self.im_run, True, "Working…")
        self.im_progress.set(0)
        self._log(self.im_log, f"\n── Starting: {len(self.im_files.files)} file(s) ──")
        threading.Thread(
            target=self._image_worker, daemon=True,
            kwargs=dict(
                files=list(self.im_files.files),
                old_bytes=self.im_old.bytes, new_bytes=self.im_new.bytes,
                match_dims=bool(self.im_dims.get()),
            ),
        ).start()

    def _image_worker(self, files, old_bytes, new_bytes, match_dims):
        total = len(files)
        grand = changed = 0
        for i, path in enumerate(files, start=1):
            name = os.path.basename(path)
            document, result = replace_images_in_document(
                path, old_bytes, new_bytes, match_by_dimensions=match_dims
            )
            if not result.ok:
                self._logq(self.im_log, f"✖ {name}: {result.error}")
            elif result.total == 0:
                self._logq(self.im_log, f"•  {name}: no matching image")
            else:
                parts = []
                if result.exact:
                    parts.append(f"{result.exact} exact")
                if result.dimensions:
                    parts.append(f"{result.dimensions} by size")
                detail = ", ".join(parts)
                if result.format_changed:
                    detail += ", format updated"
                saved = self._save_with_backup(path, document, self.im_log, name,
                                               f"{result.total} image(s) — {detail}")
                if saved:
                    grand += result.total
                    changed += 1
            self._progq(self.im_progress, i / total)
        self._logq(self.im_log,
                   f"── Done: {grand} image(s) replaced across {changed} file(s). ──")
        self._doneq(self.im_run, "Replace image")

    # ---------------------------------------------------- shared save + backup
    def _save_with_backup(self, path, document, log_box, name, summary):
        """Back up original to .bak (once) then overwrite. Returns True on
        success; logs the outcome either way."""
        try:
            bak = path + ".bak"
            if not os.path.exists(bak):
                shutil.copy2(path, bak)
                note = "backed up"
            else:
                note = ".bak already existed, kept original"
            document.save(path)
            self._logq(log_box, f"✔ {name}: {summary} — {note}")
            return True
        except PermissionError:
            self._logq(log_box, f"✖ {name}: could not save (is it open in Word?)")
        except Exception as exc:
            self._logq(log_box, f"✖ {name}: save failed — {exc}")
        return False

    # -------------------------------------------------- thread-safe UI helpers
    def _ui(self, fn):
        self._msg_queue.put(fn)

    def _logq(self, log_box, text):
        self._ui(lambda: self._log(log_box, text))

    def _progq(self, progress, value):
        self._ui(lambda: progress.set(value))

    def _doneq(self, run_btn, idle_text):
        def finish():
            self._running = False
            run_btn.configure(state="normal", text=idle_text)
        self._ui(finish)

    def _drain_queue(self):
        try:
            while True:
                self._msg_queue.get_nowait()()
        except queue.Empty:
            pass
        self.after(100, self._drain_queue)

    def _set_running(self, run_btn, running, busy_text):
        self._running = running
        run_btn.configure(state="disabled" if running else "normal",
                          text=busy_text)

    def _log(self, log_box, text):
        log_box.configure(state="normal")
        log_box.insert("end", text + "\n")
        log_box.see("end")
        log_box.configure(state="disabled")


def _selftest_pages(outfile):
    """Headless check that Word automation works in this (possibly frozen)
    build. Writes a one-line result to ``outfile``."""
    import tempfile
    from docx import Document
    try:
        p = os.path.join(tempfile.gettempdir(), "wh_selftest.docx")
        doc = Document()
        doc.add_paragraph("page one")
        doc.add_page_break()
        doc.add_paragraph("page two")
        doc.save(p)
        n = count_pages(p)
        result = f"OK pages={n}"
    except Exception as exc:
        result = f"ERR {exc}"
    with open(outfile, "w", encoding="utf-8") as fh:
        fh.write(result)


if __name__ == "__main__":
    import sys
    if len(sys.argv) >= 3 and sys.argv[1] == "--selftest-pages":
        _selftest_pages(sys.argv[2])
    else:
        WordHelperApp().mainloop()
