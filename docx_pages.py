"""Page-level operations (extract / remove / add) driven by Microsoft Word.

A .docx file has no stored page boundaries — pages are produced by Word's layout
engine. So these operations automate a hidden Word instance, which knows the real
page numbers and preserves all formatting, headers/footers, sections and images.

Requires Microsoft Word to be installed (Windows only). Every function raises
WordUnavailable if Word can't be started.

Fidelity strategy for extraction: copy the whole source file, then delete the
pages outside the wanted range. Nothing is rebuilt, so headers/footers/styles
come along untouched.
"""

from __future__ import annotations

import copy
import os
import re
import shutil
import tempfile
from dataclasses import dataclass

import pythoncom
import win32com.client
from docx import Document
from docx.oxml.ns import qn

# Word enum constants
WD_STAT_PAGES = 2
WD_GOTO_PAGE = 1
WD_GOTO_ABSOLUTE = 1
WD_COLLAPSE_END = 0
WD_SECTION_BREAK_NEXT_PAGE = 2
WD_FORMAT_DOCX = 16          # wdFormatDocumentDefault (.docx)
WD_DO_NOT_SAVE = 0


class WordUnavailable(RuntimeError):
    pass


@dataclass
class PageOpResult:
    ok: bool = True
    error: str = ""
    output: str = ""          # path written (extract/add to new file)
    pages_affected: int = 0
    source_removed: int = 0   # pages removed from source (extract w/ remove)


class WordSession:
    """Context manager for a hidden Word application. Initialises COM for the
    calling thread (needed because we run inside worker threads)."""

    def __enter__(self):
        pythoncom.CoInitialize()
        try:
            self.app = win32com.client.DispatchEx("Word.Application")
        except Exception as exc:
            pythoncom.CoUninitialize()
            raise WordUnavailable(f"Microsoft Word could not be started: {exc}")
        self.app.Visible = False
        self.app.DisplayAlerts = 0
        return self.app

    def __exit__(self, *exc):
        try:
            self.app.Quit(WD_DO_NOT_SAVE)
        except Exception:
            pass
        finally:
            pythoncom.CoUninitialize()


def _page_start(doc, page_number):
    """Character position where ``page_number`` begins."""
    return doc.GoTo(WD_GOTO_PAGE, WD_GOTO_ABSOLUTE, page_number).Start


def _strip_trailing_blanks(doc):
    """Remove trailing empty / page-break-only paragraphs left behind after a
    deletion, so removing the last pages doesn't leave a blank page."""
    for _ in range(50):
        if doc.Paragraphs.Count <= 1:
            break
        last = doc.Paragraphs.Last.Range
        # str.strip() also removes the form-feed (\x0c) of a page break.
        if last.Text.strip() == "":
            last.Delete()
        else:
            break


def _trim_to_range(doc, start, end):
    """Delete every page in ``doc`` outside [start, end] (1-based, inclusive)."""
    total = int(doc.ComputeStatistics(WD_STAT_PAGES))
    start = max(1, start)
    end = min(end, total)
    # Delete the tail first so earlier positions stay valid.
    if end < total:
        tail = _page_start(doc, end + 1)
        doc.Range(tail, doc.Content.End).Delete()
    if start > 1:
        head_end = _page_start(doc, start)
        doc.Range(0, head_end).Delete()
    return end - start + 1


def count_pages(path):
    """Return the rendered page count of ``path`` (opens Word)."""
    with WordSession() as app:
        doc = app.Documents.Open(os.path.abspath(path), False, True)  # read-only
        try:
            return int(doc.ComputeStatistics(WD_STAT_PAGES))
        finally:
            doc.Close(WD_DO_NOT_SAVE)


def extract_pages(source, dest, start, end, remove_from_source=False,
                  backup=True):
    """Write pages [start, end] of ``source`` to ``dest`` (a new .docx),
    optionally deleting those pages from ``source`` too."""
    result = PageOpResult()
    try:
        dest = os.path.abspath(dest)
        os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
        shutil.copy2(source, dest)  # full-fidelity starting point
        with WordSession() as app:
            doc = app.Documents.Open(dest)
            result.pages_affected = _trim_to_range(doc, start, end)
            doc.SaveAs2(dest, WD_FORMAT_DOCX)
            doc.Close(WD_DO_NOT_SAVE)

            if remove_from_source:
                if backup:
                    bak = source + ".bak"
                    if not os.path.exists(bak):
                        shutil.copy2(source, bak)
                sdoc = app.Documents.Open(os.path.abspath(source))
                total = int(sdoc.ComputeStatistics(WD_STAT_PAGES))
                s = max(1, start)
                e = min(end, total)
                if e < total:
                    tail = _page_start(sdoc, e + 1)
                else:
                    tail = sdoc.Content.End
                head = _page_start(sdoc, s) if s > 1 else 0
                sdoc.Range(head, tail).Delete()
                _strip_trailing_blanks(sdoc)
                sdoc.Save()
                sdoc.Close(WD_DO_NOT_SAVE)
                result.source_removed = e - s + 1
        result.output = dest
    except WordUnavailable as exc:
        result.ok = False
        result.error = str(exc)
    except Exception as exc:
        result.ok = False
        result.error = str(exc)
    return result


def remove_pages(path, start, end, backup=True):
    """Delete pages [start, end] from ``path`` in place."""
    result = PageOpResult()
    try:
        if backup:
            bak = path + ".bak"
            if not os.path.exists(bak):
                shutil.copy2(path, bak)
        with WordSession() as app:
            doc = app.Documents.Open(os.path.abspath(path))
            total = int(doc.ComputeStatistics(WD_STAT_PAGES))
            s = max(1, start)
            e = min(end, total)
            tail = _page_start(doc, e + 1) if e < total else doc.Content.End
            head = _page_start(doc, s) if s > 1 else 0
            doc.Range(head, tail).Delete()
            _strip_trailing_blanks(doc)
            doc.Save()
            doc.Close(WD_DO_NOT_SAVE)
            result.pages_affected = e - s + 1
        result.output = path
    except WordUnavailable as exc:
        result.ok = False
        result.error = str(exc)
    except Exception as exc:
        result.ok = False
        result.error = str(exc)
    return result


def add_pages(target, source, dest=None, insert_after="end",
              src_start=None, src_end=None, backup=True):
    """Insert ``source`` (optionally only pages [src_start, src_end]) into
    ``target``. ``insert_after`` is 'end' or a 1-based page number. If ``dest``
    is given the result is written there (target untouched); otherwise ``target``
    is modified in place (with a .bak backup)."""
    result = PageOpResult()
    temp = None
    try:
        insert_source = os.path.abspath(source)
        # If only a page range of the source is wanted, extract it to a temp file.
        if src_start is not None or src_end is not None:
            temp = os.path.join(tempfile.gettempdir(), "wh_addpages_tmp.docx")
            sub = extract_pages(source, temp, src_start or 1,
                                src_end or 10_000_000)
            if not sub.ok:
                return sub
            insert_source = temp

        if dest:
            dest = os.path.abspath(dest)
            os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
            shutil.copy2(target, dest)
            work = dest
        else:
            if backup:
                bak = target + ".bak"
                if not os.path.exists(bak):
                    shutil.copy2(target, bak)
            work = os.path.abspath(target)

        with WordSession() as app:
            doc = app.Documents.Open(work)
            if str(insert_after).lower() == "end":
                rng = doc.Content
                rng.Collapse(WD_COLLAPSE_END)
                rng.InsertBreak(WD_SECTION_BREAK_NEXT_PAGE)
                rng.Collapse(WD_COLLAPSE_END)
            else:
                page = int(insert_after)
                total = int(doc.ComputeStatistics(WD_STAT_PAGES))
                if page >= total:
                    rng = doc.Content
                    rng.Collapse(WD_COLLAPSE_END)
                    rng.InsertBreak(WD_SECTION_BREAK_NEXT_PAGE)
                    rng.Collapse(WD_COLLAPSE_END)
                else:
                    # Insert a section break so the inserted document keeps its
                    # own headers/footers instead of adopting the target's.
                    pos = _page_start(doc, page + 1)
                    rng = doc.Range(pos, pos)
                    rng.InsertBreak(WD_SECTION_BREAK_NEXT_PAGE)
                    rng.Collapse(WD_COLLAPSE_END)
            rng.InsertFile(insert_source)
            doc.Save()
            result.pages_affected = int(doc.ComputeStatistics(WD_STAT_PAGES))
            doc.Close(WD_DO_NOT_SAVE)
        result.output = work
    except WordUnavailable as exc:
        result.ok = False
        result.error = str(exc)
    except Exception as exc:
        result.ok = False
        result.error = str(exc)
    finally:
        if temp and os.path.exists(temp):
            try:
                os.remove(temp)
            except OSError:
                pass
    return result


# =====================================================================
# Record-based operations (a "record" / "document" == one Word section).
#
# When Word does a mail merge to a single file, each record becomes its own
# section, which is why each can have unique headers/footers. So moving whole
# records means moving whole sections — done by copying the file and trimming
# sections, which keeps every section's headers/footers/images intact.
# =====================================================================


def _section_body_groups(doc):
    """Split the body into per-section lists of paragraph text. len == #sections."""
    body = doc.element.body
    groups, cur = [], []
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            cur.append("".join(t.text or "" for t in child.findall(".//" + qn("w:t"))))
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                groups.append(cur)
                cur = []
    groups.append(cur)  # final (body-level) section
    return groups


def _trailing_empty_count(doc):
    """How many trailing sections have no text (Word mail-merge leaves one)."""
    n = 0
    for g in reversed(_section_body_groups(doc)):
        if any(t.strip() for t in g):
            break
        n += 1
    return n


def count_records(path):
    """Number of real records (sections) in ``path``, excluding the empty
    trailing section that Word's mail merge leaves behind. Uses python-docx."""
    doc = Document(path)
    return max(0, len(doc.sections) - _trailing_empty_count(doc))


def _strip_trailing_empty_section(path):
    """If the document ends with an empty section, remove it (promoting the last
    real section's properties to the body level). No-op otherwise."""
    doc = Document(path)
    if _trailing_empty_count(doc) == 0:
        return
    _promote_last_section(path)


def _promote_last_section(path):
    """After trimming trailing sections, the removed section's properties can
    linger at the document-body level, leaving a phantom empty section. Promote
    the last *kept* section's properties to the body level and drop the leftover.
    """
    d = Document(path)
    body = d.element.body
    target = None
    for p in body.findall(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
            target = (p, pPr, pPr.find(qn("w:sectPr")))
    if target is None:
        return
    p_elem, pPr_elem, sp_elem = target
    # Delete any paragraphs that belong to the leftover trailing section.
    seen = False
    for p in list(body.findall(qn("w:p"))):
        if seen:
            body.remove(p)
        if p is p_elem:
            seen = True
    body_sectPr = body.find(qn("w:sectPr"))
    new_sectPr = copy.deepcopy(sp_elem)
    if body_sectPr is not None:
        body.remove(body_sectPr)
    body.append(new_sectPr)
    pPr_elem.remove(sp_elem)
    d.save(path)


_DATE_RE = re.compile(
    r"^\s*(January|February|March|April|May|June|July|August|September|"
    r"October|November|December)\s+\d{1,2},?\s+\d{4}\s*$", re.I)


def record_names(path):
    """A short label for each real record — the recipient line of each section
    (first non-empty, non-date body line). Aligns 1:1 with count_records()."""
    doc = Document(path)
    groups = _section_body_groups(doc)
    te = _trailing_empty_count(doc)
    real = groups[: len(groups) - te] if te else groups
    names = []
    for i, g in enumerate(real, start=1):
        lines = [t.strip() for t in g if t.strip()]
        name = next((ln for ln in lines if not _DATE_RE.match(ln)), None)
        if not name and lines:
            name = lines[0]
        names.append((name or f"Document {i}")[:60])
    return names


def _group_ranges(nums):
    """Group a list of ints into contiguous [start, end] ranges."""
    out = []
    for n in sorted(set(nums)):
        if out and n == out[-1][1] + 1:
            out[-1][1] = n
        else:
            out.append([n, n])
    return out


def _delete_sections(doc, del_indices):
    """Delete the given 1-based section indices from an open Word document.
    Returns True if the document's tail (last section) was deleted."""
    total = doc.Sections.Count
    del_indices = [k for k in del_indices if 1 <= k <= total]
    if not del_indices:
        return False
    # Capture absolute start positions up front; deleting higher-numbered
    # sections never shifts the positions of lower-numbered content.
    starts = {k: doc.Sections(k).Range.Start for k in range(1, total + 1)}
    content_end = doc.Content.End
    tail = False
    for a, b in sorted(_group_ranges(del_indices), reverse=True):
        start_pos = starts[a] if a > 1 else 0
        if b < total:
            end_pos = starts[b + 1]
        else:
            end_pos = content_end
            tail = True
        doc.Range(start_pos, end_pos).Delete()
    return tail


def _normalize_indices(indices, total):
    return sorted({int(i) for i in indices if 1 <= int(i) <= total})


def extract_records(source, dest, indices, remove_from_source=False, backup=True):
    """Write records ``indices`` (a list of 1-based record numbers, any order,
    possibly non-contiguous) of ``source`` to ``dest`` with full fidelity
    (headers/footers/images preserved). Optionally remove them from the source."""
    result = PageOpResult()
    try:
        total = count_records(source)
        indices = _normalize_indices(indices, total)
        if not indices:
            result.ok = False
            result.error = f"no valid documents selected (source has {total})"
            return result
        dest = os.path.abspath(dest)
        os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
        shutil.copy2(source, dest)
        keep = set(indices)
        with WordSession() as app:
            doc = app.Documents.Open(dest)
            sec_total = doc.Sections.Count
            del_idx = [k for k in range(1, sec_total + 1) if k not in keep]
            tail = _delete_sections(doc, del_idx)
            doc.SaveAs2(dest, WD_FORMAT_DOCX)
            doc.Close(WD_DO_NOT_SAVE)
        if tail:
            _promote_last_section(dest)
        _strip_trailing_empty_section(dest)
        result.output = dest
        result.pages_affected = len(indices)

        if remove_from_source:
            sub = remove_records(source, indices, backup=backup)
            if not sub.ok:
                result.ok = False
                result.error = f"extracted OK, but removing from source failed: {sub.error}"
            else:
                result.source_removed = sub.pages_affected
    except WordUnavailable as exc:
        result.ok = False
        result.error = str(exc)
    except Exception as exc:
        result.ok = False
        result.error = str(exc)
    return result


def remove_records(path, indices, backup=True):
    """Delete records ``indices`` (1-based, any order) from ``path`` in place."""
    result = PageOpResult()
    try:
        total = count_records(path)
        indices = _normalize_indices(indices, total)
        if not indices:
            result.ok = False
            result.error = f"no valid documents selected (file has {total})"
            return result
        if backup:
            bak = path + ".bak"
            if not os.path.exists(bak):
                shutil.copy2(path, bak)
        with WordSession() as app:
            doc = app.Documents.Open(os.path.abspath(path))
            tail = _delete_sections(doc, indices)
            doc.Save()
            doc.Close(WD_DO_NOT_SAVE)
        if tail:
            _promote_last_section(path)
        result.pages_affected = len(indices)
        result.output = path
    except WordUnavailable as exc:
        result.ok = False
        result.error = str(exc)
    except Exception as exc:
        result.ok = False
        result.error = str(exc)
    return result


def _append_document(work_path, insert_path):
    """Append ``insert_path`` to the end of the already-open-able ``work_path``
    as its own section(s), preserving the inserted file's headers/footers."""
    with WordSession() as app:
        doc = app.Documents.Open(os.path.abspath(work_path))
        rng = doc.Content
        rng.Collapse(WD_COLLAPSE_END)
        rng.InsertBreak(WD_SECTION_BREAK_NEXT_PAGE)
        rng.Collapse(WD_COLLAPSE_END)
        rng.InsertFile(os.path.abspath(insert_path))
        doc.Save()
        pages = int(doc.ComputeStatistics(WD_STAT_PAGES))
        doc.Close(WD_DO_NOT_SAVE)
    return pages


def move_records(source, target, indices, dest=None, backup=True):
    """Move records ``indices`` (1-based, any order, non-contiguous OK) out of
    ``source`` and append them to the end of ``target`` (preserving each
    record's headers/footers). If ``dest`` is given the combined result is
    written there and ``target`` is left untouched; otherwise ``target`` is
    modified in place (.bak backup). The records are removed from ``source``
    (.bak backup)."""
    result = PageOpResult()
    temp = os.path.join(tempfile.gettempdir(), "wh_move_tmp.docx")
    try:
        ext = extract_records(source, temp, indices, remove_from_source=False)
        if not ext.ok:
            return ext

        if dest:
            dest = os.path.abspath(dest)
            os.makedirs(os.path.dirname(dest) or ".", exist_ok=True)
            shutil.copy2(target, dest)
            work = dest
        else:
            if backup:
                bak = target + ".bak"
                if not os.path.exists(bak):
                    shutil.copy2(target, bak)
            work = os.path.abspath(target)

        # Drop the target's empty trailing section so the moved records append
        # right after the last real record (not after a stranded blank section).
        _strip_trailing_empty_section(work)
        result.pages_affected = _append_document(work, temp)
        result.output = work

        rem = remove_records(source, indices, backup=backup)
        if not rem.ok:
            result.ok = False
            result.error = f"added to target, but removing from source failed: {rem.error}"
        else:
            result.source_removed = rem.pages_affected
    except WordUnavailable as exc:
        result.ok = False
        result.error = str(exc)
    except Exception as exc:
        result.ok = False
        result.error = str(exc)
    finally:
        if os.path.exists(temp):
            try:
                os.remove(temp)
            except OSError:
                pass
    return result


def word_available():
    """Quick check whether Word can be automated."""
    try:
        with WordSession():
            return True
    except Exception:
        return False
